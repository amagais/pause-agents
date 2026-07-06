"""Export ICU-PAUSE output to Word (.docx), PDF, and EMR-ready plain text.

Design principles (Apple HIG-inspired):
- Typography hierarchy over color to create structure
- Generous whitespace — let content breathe
- Color reserved for things that need attention (alerts, low-confidence)
- Thin separators instead of heavy colored bars
- Clean, scannable layout optimized for print and screen
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from icu_pause.rendering.formatter import (
    SECTION_ORDER,
    _format_clinician_warning_line,
    render_warnings_for_clinician,
)
from icu_pause.rendering.citations import (
    line_plain_prefix,
    render_section,
    segments_by_line,
)

# Confidence threshold below which a section is flagged for review
LOW_CONFIDENCE_THRESHOLD = 0.5

# ---------------------------------------------------------------------------
# Color palette — restrained, semantic use only
# ---------------------------------------------------------------------------
_ACCENT = (0, 122, 255)          # SF Blue — accent color for section letters
_TEXT_PRIMARY = (28, 28, 30)     # Near-black — primary text
_TEXT_SECONDARY = (99, 99, 102)  # Gray — metadata, timestamps
_TEXT_TERTIARY = (142, 142, 147) # Light gray — footer, hints
_SEPARATOR = (209, 209, 214)     # System gray 4 — thin rules
_BG_WHITE = (255, 255, 255)

# Alert colors — only used when something needs attention
_QA_ACCENT = (255, 59, 48)      # SF Red
_QA_BG = (255, 243, 242)
_QA_TEXT = (153, 27, 27)
_WARN_ACCENT = (255, 149, 0)    # SF Orange
_WARN_BG = (255, 247, 237)
_WARN_TEXT = (120, 53, 15)
_LOW_CONF_ACCENT = (255, 149, 0)
_LOW_CONF_BG = (255, 247, 237)
_LOW_CONF_TEXT = (67, 20, 7)
_UNCERT_ACCENT = (255, 204, 0)  # SF Yellow
_UNCERT_BG = (255, 251, 235)
_UNCERT_TEXT = (120, 53, 15)

# Problem header accent — subtle teal
_PROBLEM_ACCENT = (0, 122, 255)

# Citation tier colors — matches frontend/src/App.css and review_app theme.
# Muted/darker here because print contrast is less forgiving.
_CITE_DECISION_CRITICAL = (30, 64, 175)  # deep blue
_CITE_UNVERIFIED = (180, 83, 9)          # amber


# ---------------------------------------------------------------------------
# Word (.docx) export
# ---------------------------------------------------------------------------

def export_docx(output: dict[str, Any]) -> BytesIO:
    """Generate a styled Word document from ICU-PAUSE output."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # -- Page margins (generous) --
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Default paragraph style --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(*_TEXT_PRIMARY)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.2

    # ── Title block ──────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(2)
    run = p_title.add_run("ICU to Ward Transfer Summary")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*_TEXT_PRIMARY)
    run.font.name = "Calibri"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_after = Pt(4)
    run = p_sub.add_run("ICU-PAUSE Framework")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*_TEXT_SECONDARY)
    run.font.name = "Calibri"

    # Thin accent line under title
    _add_separator(doc, _ACCENT, weight=8)

    # ── Metadata line ────────────────────────────────────────────
    hosp_id = output.get("hospitalization_id", "N/A")
    generated = output.get("generated_at", "N/A")
    if isinstance(generated, str) and "T" in generated:
        generated = generated.split("T")[0]  # Just the date

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"Patient ID: {hosp_id}    \u00b7    {generated}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*_TEXT_SECONDARY)
    run.font.name = "Calibri"

    # ── Alert banners ────────────────────────────────────────────
    qa_issues = output.get("qa_issues", [])
    if qa_issues:
        _add_alert_banner_docx(
            doc, "QA Issues \u2014 Review Required",
            qa_issues, _QA_BG, _QA_ACCENT, _QA_TEXT,
        )

    visible_warnings = render_warnings_for_clinician(output.get("warnings", []))
    if visible_warnings:
        _add_alert_banner_docx(
            doc, "Warnings",
            [_format_clinician_warning_line(w) for w in visible_warnings],
            _WARN_BG, _WARN_ACCENT, _WARN_TEXT,
        )

    # ── ICU-PAUSE Sections ───────────────────────────────────────
    sections = output.get("sections", {})
    confidences = output.get("section_confidences", {})
    citation_index = output.get("metadata", {}).get("citation_index", {}) or {}

    for key, letter, label in SECTION_ORDER:
        content = sections.get(key, "Not enough information from structured data.")
        confidence = confidences.get(key, 1.0)
        is_uncertainty = key == "U_uncertainty"
        is_low_conf = confidence < LOW_CONFIDENCE_THRESHOLD
        needs_review = is_uncertainty or is_low_conf

        # Determine colors
        if is_low_conf:
            accent_color = _LOW_CONF_ACCENT
            body_bg = _LOW_CONF_BG
            body_text = _LOW_CONF_TEXT
        elif is_uncertainty:
            accent_color = _UNCERT_ACCENT
            body_bg = _UNCERT_BG
            body_text = _UNCERT_TEXT
        else:
            accent_color = _ACCENT
            body_bg = None  # no background for normal sections
            body_text = _TEXT_PRIMARY

        # Section header: letter in accent color + label in dark text
        p_header = doc.add_paragraph()
        p_header.paragraph_format.space_before = Pt(14)
        p_header.paragraph_format.space_after = Pt(4)

        run_letter = p_header.add_run(f"{letter}  ")
        run_letter.font.size = Pt(14)
        run_letter.font.bold = True
        run_letter.font.color.rgb = RGBColor(*accent_color)
        run_letter.font.name = "Calibri"

        run_label = p_header.add_run(label)
        run_label.font.size = Pt(12)
        run_label.font.bold = True
        run_label.font.color.rgb = RGBColor(*_TEXT_PRIMARY)
        run_label.font.name = "Calibri"

        # Badge for review items
        if needs_review:
            badge_parts = []
            if is_low_conf:
                badge_parts.append(f"Low Confidence ({confidence:.0%})")
            if is_uncertainty:
                badge_parts.append("Review Recommended")
            sep = " \u00b7 "
            badge_run = p_header.add_run(f"    {sep.join(badge_parts)}")
            badge_run.font.size = Pt(9)
            badge_run.font.color.rgb = RGBColor(*accent_color)
            badge_run.font.name = "Calibri"

        # Thin separator under header
        _add_separator(doc, _SEPARATOR, weight=4)

        # Section body — run the canonical citation renderer once per
        # section so numbering resets at the section boundary (same as
        # the Streamlit/React renderers).  The resulting footnote list
        # is emitted at the end of the body in small grey text, since
        # hover tooltips don't work on print.
        rendered = render_section(content, citation_index)
        if body_bg and body_bg != _BG_WHITE:
            # Use a single-cell table for colored background
            card = doc.add_table(rows=1, cols=1)
            card.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(card, body_bg)  # subtle matching border
            cell = card.cell(0, 0)
            _set_cell_shading(cell, body_bg)
            _set_cell_margins(cell, top=120, bottom=120, left=200, right=200)
            _add_rich_content(cell, rendered["segments"], body_text)
            _append_footnotes_to_cell(cell, rendered["footnotes"])
        else:
            # No background — add directly as paragraphs
            _add_rich_content_direct(doc, rendered["segments"], body_text)
            _append_footnotes_direct(doc, rendered["footnotes"])

        # Spacer
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(2)

    # ── To-Do Checklist ──────────────────────────────────────────
    todo_items = output.get("todo_checklist", [])
    if todo_items:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("To-Do List")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*_TEXT_PRIMARY)
        run.font.name = "Calibri"

        _add_separator(doc, _SEPARATOR, weight=4)

        from icu_pause.rendering.formatter import _BUCKET_LABELS, _BUCKET_ORDER
        normalized = [
            t if isinstance(t, dict) else {"bucket": "ward_ongoing", "text": t}
            for t in todo_items
        ]
        for bkey in _BUCKET_ORDER:
            items = [t["text"] for t in normalized if t.get("bucket", "ward_ongoing") == bkey]
            if not items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(_BUCKET_LABELS[bkey])
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*_TEXT_SECONDARY)
            run.font.name = "Calibri"
            for text in items:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12)
                run = p.add_run(f"\u25cb  {text}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(*_TEXT_PRIMARY)
                run.font.name = "Calibri"

    # ── Footer ───────────────────────────────────────────────────
    _add_separator(doc, _SEPARATOR, weight=4)
    meta = output.get("metadata", {})
    if meta:
        filled = meta.get("sections_filled", "?")
        total = meta.get("sections_total", "?")
        agents = meta.get("agent_count", "?")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(f"{filled}/{total} sections  \u00b7  {agents} agents  \u00b7  AI-generated, physician review required")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(*_TEXT_TERTIARY)
        run.font.name = "Calibri"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _add_separator(doc, rgb: tuple[int, int, int], weight: int = 6) -> None:
    """Add a thin horizontal rule as a table with colored top border."""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Only top border, rest none
    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
      <w:top w:val="single" w:sz="{weight}" w:space="0" w:color="{hex_color}"/>
      <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>
    """
    tblPr = table._tbl.tblPr
    if tblPr is None:
        from lxml import etree
        tblPr = etree.SubElement(table._tbl, qn("w:tblPr"))
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))
    # Full width
    width_xml = f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(parse_xml(width_xml))

    # Make the cell empty and tiny
    cell = table.cell(0, 0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    _set_cell_margins(cell, top=0, bottom=0, left=0, right=0)


def _set_cell_shading(cell, rgb: tuple[int, int, int]) -> None:
    """Set background color of a table cell."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0) -> None:
    """Set cell margins in twips (1/20 of a point)."""
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = tcPr.find(qn("w:tcMar"))
    if margins is None:
        from lxml import etree
        margins = etree.SubElement(tcPr, qn("w:tcMar"))
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        elem = margins.find(qn(f"w:{side}"))
        if elem is None:
            from lxml import etree
            elem = etree.SubElement(margins, qn(f"w:{side}"))
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")


def _set_table_borders(table, rgb: tuple[int, int, int]) -> None:
    """Set table borders to a single color."""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>
      <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>
    """
    tblPr = table._tbl.tblPr
    if tblPr is None:
        from lxml import etree
        tblPr = etree.SubElement(table._tbl, qn("w:tblPr"))
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))

    width_xml = f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(parse_xml(width_xml))


def _add_alert_banner_docx(
    doc, title: str, items: list[str],
    bg: tuple, accent: tuple, text_color: tuple,
) -> None:
    """Add a colored alert banner with left accent border."""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left border accent only
    hex_accent = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"
    hex_bg = f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}"
    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
      <w:top w:val="single" w:sz="2" w:space="0" w:color="{hex_bg}"/>
      <w:left w:val="single" w:sz="18" w:space="0" w:color="{hex_accent}"/>
      <w:bottom w:val="single" w:sz="2" w:space="0" w:color="{hex_bg}"/>
      <w:right w:val="single" w:sz="2" w:space="0" w:color="{hex_bg}"/>
      <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>
    """
    tblPr = table._tbl.tblPr
    if tblPr is None:
        from lxml import etree
        tblPr = etree.SubElement(table._tbl, qn("w:tblPr"))
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))
    width_xml = f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(parse_xml(width_xml))

    cell = table.cell(0, 0)
    _set_cell_shading(cell, bg)
    _set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*text_color)
    run.font.name = "Calibri"

    for item in items:
        p = cell.add_paragraph()
        run = p.add_run(f"\u2022  {item}")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(*text_color)
        run.font.name = "Calibri"

    # Spacer after banner
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)


def _add_cite_segments_to_paragraph(
    p, line_segments: list[dict[str, Any]],
    body_text: tuple[int, int, int],
    prefix_text: str = "",
    font_size_pt: float = 10.5,
    bold_text: bool = False,
) -> None:
    """Emit *line_segments* as runs on paragraph *p*.

    Text segments become regular runs (colored with *body_text*).
    Citation segments become superscript runs colored per tier — this is
    the DOCX equivalent of the ``.icp-cite`` CSS class in the web UIs.
    """
    from docx.shared import Pt, RGBColor

    if prefix_text:
        run = p.add_run(prefix_text)
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = RGBColor(*body_text)
        run.font.name = "Calibri"
        run.font.bold = bold_text

    for seg in line_segments:
        if seg["kind"] == "text":
            run = p.add_run(seg["text"])
            run.font.size = Pt(font_size_pt)
            run.font.color.rgb = RGBColor(*body_text)
            run.font.name = "Calibri"
            run.font.bold = bold_text
        elif seg["kind"] == "cite":
            # Embed the superscript inline.  We use Unicode superscript
            # characters (Calibri handles them) rather than
            # run.font.superscript=True so the character sizing stays
            # consistent regardless of the reader's Word version.
            run = p.add_run(seg["marker"])
            run.font.size = Pt(font_size_pt)
            run.font.bold = True
            color = (
                _CITE_UNVERIFIED if seg["tier"] == "unverified"
                else _CITE_DECISION_CRITICAL
            )
            run.font.color.rgb = RGBColor(*color)
            run.font.name = "Calibri"


def _render_segment_line_docx(
    p, line_segments: list[dict[str, Any]], text_color: tuple[int, int, int],
) -> None:
    """Apply structural detection (checkbox / header / bullet) to a line.

    Routes to the appropriate paragraph formatting, then calls
    ``_add_cite_segments_to_paragraph`` to render the line body including
    any inline citations.
    """
    from docx.shared import Pt

    prefix = line_plain_prefix(line_segments)

    if re.match(r'^\[[ x]?\]\s', prefix):
        # Strip "[] " from the FIRST text segment so the label text starts
        # cleanly; citations embedded later in the line are preserved.
        stripped = _strip_first_text(line_segments, r'^\[[ x]?\]\s*')
        p.paragraph_format.left_indent = Pt(8)
        _add_cite_segments_to_paragraph(p, stripped, text_color, prefix_text="\u25cb  ")
        return

    if prefix.startswith("## "):
        # Domain subheading — no citations rendered (they're rare here)
        stripped = _strip_first_text(line_segments, r'^##\s+')
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        # Uppercase the text-segment parts only.
        for seg in stripped:
            if seg["kind"] == "text":
                seg["text"] = seg["text"].upper()
        _add_cite_segments_to_paragraph(
            p, stripped, _TEXT_SECONDARY, font_size_pt=9, bold_text=True,
        )
        return

    if prefix.startswith("#"):
        stripped = _strip_first_text(line_segments, r'^#+\s*')
        p.paragraph_format.left_indent = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        _add_cite_segments_to_paragraph(
            p, stripped, _PROBLEM_ACCENT, bold_text=True,
        )
        return

    if prefix.startswith("- "):
        stripped = _strip_first_text(line_segments, r'^-\s+')
        p.paragraph_format.left_indent = Pt(16)
        _add_cite_segments_to_paragraph(
            p, stripped, _TEXT_TERTIARY, font_size_pt=9,
        )
        return

    _add_cite_segments_to_paragraph(p, line_segments, text_color)


def _strip_first_text(
    line_segments: list[dict[str, Any]], pattern: str,
) -> list[dict[str, Any]]:
    """Return a copy of *line_segments* with *pattern* removed from the first text segment."""
    out: list[dict[str, Any]] = []
    stripped = False
    for seg in line_segments:
        if not stripped and seg["kind"] == "text":
            out.append({"kind": "text", "text": re.sub(pattern, "", seg["text"])})
            stripped = True
        else:
            out.append(dict(seg))
    return out


def _add_rich_content(
    cell, segments: list[dict[str, Any]], text_color: tuple[int, int, int],
) -> None:
    """Render pre-computed citation segments into a table cell."""
    lines = segments_by_line(segments)
    first_line = True

    for line_segments in lines:
        prefix = line_plain_prefix(line_segments)
        if not prefix and not any(s["kind"] == "cite" for s in line_segments):
            continue

        p = cell.paragraphs[0] if first_line else cell.add_paragraph()
        first_line = False
        _render_segment_line_docx(p, line_segments, text_color)


def _add_rich_content_direct(
    doc, segments: list[dict[str, Any]], text_color: tuple[int, int, int],
) -> None:
    """Render pre-computed citation segments directly into the document."""
    from docx.shared import Pt

    lines = segments_by_line(segments)
    for line_segments in lines:
        prefix = line_plain_prefix(line_segments)
        if not prefix and not any(s["kind"] == "cite" for s in line_segments):
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _render_segment_line_docx(p, line_segments, text_color)


def _append_footnotes_direct(doc, footnotes: list[dict[str, Any]]) -> None:
    """Emit end-of-section footnote lines so printed output is still usable."""
    from docx.shared import Pt, RGBColor

    if not footnotes:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Pt(8)
    for i, fn in enumerate(footnotes):
        if i > 0:
            run = p.add_run("   ")
            run.font.size = Pt(8)
        run = p.add_run(f"{fn['marker']} {fn['tooltip']}")
        run.font.size = Pt(8)
        color = (
            _CITE_UNVERIFIED if fn["tier"] == "unverified"
            else _TEXT_TERTIARY
        )
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"


def _append_footnotes_to_cell(cell, footnotes: list[dict[str, Any]]) -> None:
    """Cell-scoped variant of ``_append_footnotes_direct``."""
    from docx.shared import Pt, RGBColor

    if not footnotes:
        return
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    for i, fn in enumerate(footnotes):
        if i > 0:
            run = p.add_run("   ")
            run.font.size = Pt(8)
        run = p.add_run(f"{fn['marker']} {fn['tooltip']}")
        run.font.size = Pt(8)
        color = (
            _CITE_UNVERIFIED if fn["tier"] == "unverified"
            else _TEXT_TERTIARY
        )
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def _pdf_safe(text: str) -> str:
    """Replace unicode characters that fpdf's built-in Helvetica can't render."""
    return (
        text
        .replace("\u2014", " - ")   # em dash
        .replace("\u2013", "-")      # en dash
        .replace("\u00b7", "-")      # middle dot
        .replace("\u25cb", "o")      # circle → o
        .replace("\u2610", "[ ]")    # ballot box
        .replace("\u2022", "-")      # bullet
        .replace("\u2502", "|")      # box drawing
    )


def export_pdf(output: dict[str, Any]) -> BytesIO:
    """Generate a styled PDF from ICU-PAUSE output."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    # ── Title ────────────────────────────────────────────────────
    pdf.set_text_color(*_TEXT_PRIMARY)
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(page_w, 10, "ICU to Ward Transfer Summary", new_x="LMARGIN", new_y="NEXT")

    pdf.set_text_color(*_TEXT_SECONDARY)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(page_w, 6, "ICU-PAUSE Framework", new_x="LMARGIN", new_y="NEXT")

    # Accent line
    pdf.ln(2)
    pdf.set_draw_color(*_ACCENT)
    pdf.set_line_width(0.8)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + page_w, pdf.get_y())
    pdf.set_line_width(0.2)
    pdf.ln(4)

    # ── Metadata ─────────────────────────────────────────────────
    hosp_id = output.get("hospitalization_id", "N/A")
    generated = output.get("generated_at", "N/A")
    if isinstance(generated, str) and "T" in generated:
        generated = generated.split("T")[0]

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*_TEXT_SECONDARY)
    pdf.cell(page_w, 5, _pdf_safe(f"Patient ID: {hosp_id}    \u00b7    {generated}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # ── Alert banners ────────────────────────────────────────────
    qa_issues = output.get("qa_issues", [])
    if qa_issues:
        _pdf_alert_banner(pdf, "QA Issues \u2014 Review Required", qa_issues,
                          _QA_BG, _QA_ACCENT, _QA_TEXT, page_w)

    visible_warnings = render_warnings_for_clinician(output.get("warnings", []))
    if visible_warnings:
        _pdf_alert_banner(
            pdf, "Warnings",
            [_format_clinician_warning_line(w) for w in visible_warnings],
            _WARN_BG, _WARN_ACCENT, _WARN_TEXT, page_w,
        )

    # ── Sections ─────────────────────────────────────────────────
    sections = output.get("sections", {})
    confidences = output.get("section_confidences", {})
    citation_index = output.get("metadata", {}).get("citation_index", {}) or {}

    for key, letter, label in SECTION_ORDER:
        content = sections.get(key, "Not enough information from structured data.")
        confidence = confidences.get(key, 1.0)
        is_uncertainty = key == "U_uncertainty"
        is_low_conf = confidence < LOW_CONFIDENCE_THRESHOLD
        needs_review = is_uncertainty or is_low_conf

        if is_low_conf:
            accent = _LOW_CONF_ACCENT
            body_bg = _LOW_CONF_BG
            body_text = _LOW_CONF_TEXT
        elif is_uncertainty:
            accent = _UNCERT_ACCENT
            body_bg = _UNCERT_BG
            body_text = _UNCERT_TEXT
        else:
            accent = _ACCENT
            body_bg = _BG_WHITE
            body_text = _TEXT_PRIMARY

        # Section header
        pdf.ln(4)
        pdf.set_text_color(*accent)
        pdf.set_font("Helvetica", "B", 13)
        x_start = pdf.get_x()
        pdf.cell(0, 7, f"{letter}", new_x="END")

        pdf.set_text_color(*_TEXT_PRIMARY)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, f"  {label}", new_x="LMARGIN", new_y="NEXT")

        # Badge
        if needs_review:
            badge_parts = []
            if is_low_conf:
                badge_parts.append(f"Low Confidence ({confidence:.0%})")
            if is_uncertainty:
                badge_parts.append("Review Recommended")
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(*accent)
            pdf.cell(page_w, 4, _pdf_safe(" \u00b7 ".join(badge_parts)), new_x="LMARGIN", new_y="NEXT")

        # Thin separator
        pdf.set_draw_color(*_SEPARATOR)
        pdf.set_line_width(0.3)
        pdf.line(pdf.l_margin, pdf.get_y() + 1, pdf.l_margin + page_w, pdf.get_y() + 1)
        pdf.ln(3)

        # Body — run the canonical citation renderer once per section
        # so numbering resets at the section boundary (matches the web
        # renderers).  Footnotes are emitted at the end of the section;
        # Helvetica built-in can't render Unicode superscripts so the
        # inline marker format for PDF is ``[N]`` instead of ``¹²³``.
        rendered = render_section(content, citation_index)
        _pdf_rich_content(pdf, rendered["segments"], body_bg, body_text, page_w)
        _pdf_append_footnotes(pdf, rendered["footnotes"], page_w)
        pdf.ln(2)

    # ── To-Do Checklist ──────────────────────────────────────────
    todo_items = output.get("todo_checklist", [])
    if todo_items:
        from icu_pause.rendering.formatter import _BUCKET_LABELS, _BUCKET_ORDER
        normalized = [
            t if isinstance(t, dict) else {"bucket": "ward_ongoing", "text": t}
            for t in todo_items
        ]

        pdf.ln(4)
        pdf.set_text_color(*_TEXT_PRIMARY)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(page_w, 7, "To-Do List", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(*_SEPARATOR)
        pdf.set_line_width(0.3)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + page_w, pdf.get_y())
        pdf.ln(3)

        for bkey in _BUCKET_ORDER:
            items = [t["text"] for t in normalized if t.get("bucket", "ward_ongoing") == bkey]
            if not items:
                continue
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*_TEXT_SECONDARY)
            pdf.cell(page_w, 7, _pdf_safe(f"  {_BUCKET_LABELS[bkey]}"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*_TEXT_PRIMARY)
            for text in items:
                pdf.multi_cell(page_w, 5, _pdf_safe(f"      o  {text}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # ── Footer ───────────────────────────────────────────────────
    pdf.set_draw_color(*_SEPARATOR)
    pdf.set_line_width(0.3)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + page_w, pdf.get_y())
    pdf.ln(3)

    meta = output.get("metadata", {})
    if meta:
        filled = meta.get("sections_filled", "?")
        total = meta.get("sections_total", "?")
        agents = meta.get("agent_count", "?")
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*_TEXT_TERTIARY)
        pdf.cell(page_w, 5,
                 _pdf_safe(f"{filled}/{total} sections  \u00b7  {agents} agents  \u00b7  AI-generated, physician review required"),
                 new_x="LMARGIN", new_y="NEXT")

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _pdf_alert_banner(
    pdf, title: str, items: list[str],
    bg: tuple, accent: tuple, text_color: tuple,
    page_w: float,
) -> None:
    """Add a colored alert banner with left accent line."""
    y_start = pdf.get_y()

    pdf.set_fill_color(*bg)
    pdf.set_text_color(*text_color)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(page_w, 6, _pdf_safe(f"    {title}"), new_x="LMARGIN", new_y="NEXT", fill=True)

    pdf.set_font("Helvetica", "", 8.5)
    for item in items:
        pdf.multi_cell(page_w, 5, _pdf_safe(f"      \u2022  {item}"), new_x="LMARGIN", new_y="NEXT", fill=True)

    y_end = pdf.get_y()

    # Left accent border
    pdf.set_draw_color(*accent)
    pdf.set_line_width(1.2)
    pdf.line(pdf.l_margin, y_start, pdf.l_margin, y_end)
    pdf.set_line_width(0.2)
    pdf.ln(4)


def _line_text_with_brackets(line_segments: list[dict[str, Any]]) -> str:
    """Flatten a line's segments to a printable string for the PDF renderer.

    Citations render as ``[N]`` bracket markers since Helvetica's built-in
    Unicode coverage doesn't include superscript digits.  The end-of-section
    footnote list resolves the numbers back to source values.
    """
    parts: list[str] = []
    for seg in line_segments:
        if seg["kind"] == "text":
            parts.append(seg["text"])
        elif seg["kind"] == "cite":
            parts.append(f"[{seg['number']}]")
    return "".join(parts)


def _pdf_rich_content(
    pdf, segments: list[dict[str, Any]], bg: tuple, text_color: tuple, page_w: float,
) -> None:
    """Render pre-computed citation segments into PDF."""
    lines = segments_by_line(segments)
    for line_segments in lines:
        prefix = line_plain_prefix(line_segments)
        if not prefix and not any(s["kind"] == "cite" for s in line_segments):
            continue

        line = _line_text_with_brackets(line_segments)
        trimmed = line.strip()

        if re.match(r'^\[[ x]?\]\s', trimmed):
            label = re.sub(r'^\[[ x]?\]\s*', '', trimmed)
            pdf.set_text_color(*text_color)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(page_w, 5, _pdf_safe(f"      o  {label}"), new_x="LMARGIN", new_y="NEXT")

        elif prefix.startswith("## "):
            # Domain subheading
            body = re.sub(r'^##\s+', '', trimmed)
            pdf.ln(3)
            pdf.set_text_color(*_TEXT_SECONDARY)
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.cell(page_w, 5, _pdf_safe(f"  {body.upper()}"), new_x="LMARGIN", new_y="NEXT")

        elif prefix.startswith("#"):
            body = re.sub(r'^#+\s*', '', trimmed)
            pdf.set_text_color(*_PROBLEM_ACCENT)
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(page_w, 5.5, _pdf_safe(f"    {body}"), new_x="LMARGIN", new_y="NEXT")

        elif prefix.startswith("- "):
            body = re.sub(r'^-\s+', '', trimmed)
            pdf.set_text_color(*_TEXT_TERTIARY)
            pdf.set_font("Helvetica", "", 8.5)
            pdf.multi_cell(page_w, 4.5, _pdf_safe(f"          {body}"), new_x="LMARGIN", new_y="NEXT")

        else:
            pdf.set_text_color(*text_color)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(page_w, 5, _pdf_safe(f"  {line}"), new_x="LMARGIN", new_y="NEXT")


def _pdf_append_footnotes(pdf, footnotes: list[dict[str, Any]], page_w: float) -> None:
    """Append a small-text footnote block after a section's body."""
    if not footnotes:
        return
    pdf.ln(1)
    pdf.set_font("Helvetica", "", 7.5)
    for fn in footnotes:
        color = (
            _CITE_UNVERIFIED if fn["tier"] == "unverified"
            else _TEXT_TERTIARY
        )
        pdf.set_text_color(*color)
        pdf.multi_cell(
            page_w, 3.5,
            _pdf_safe(f"    [{fn['number']}] {fn['tooltip']}"),
            new_x="LMARGIN", new_y="NEXT",
        )


# ---------------------------------------------------------------------------
# EMR plain-text export
# ---------------------------------------------------------------------------

def render_emr_text(output: dict[str, Any]) -> str:
    """Render ICU-PAUSE output as plain text optimized for EMR copy-paste.

    EMRs (Epic, Cerner) accept plain text in note fields. This format:
    - Uses simple ASCII section headers (no unicode decorations)
    - Preserves checkbox format as [ ] that some EMRs recognize
    - Keeps line lengths reasonable for EMR text fields
    - Omits metadata/QA issues (internal, not patient-facing)
    - Includes a header identifying this as AI-generated
    """
    lines: list[str] = []

    lines.append("*** ICU to Ward Transfer Summary (ICU-PAUSE) ***")
    lines.append(f"Patient ID: {output.get('hospitalization_id', 'N/A')}")
    generated = output.get("generated_at", "N/A")
    if isinstance(generated, str) and "T" in generated:
        generated = generated.split("T")[0]
    lines.append(f"Date: {generated}")
    lines.append("")

    sections = output.get("sections", {})
    for key, letter, label in SECTION_ORDER:
        content = sections.get(key, "")
        if not content or content == "Not enough information from structured data.":
            continue

        lines.append(f"{letter}. {label}")
        lines.append("-" * 40)

        # Clean up content for EMR
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # Domain subheadings (## Nursing, ## Respiratory, etc.)
            if stripped.startswith("## "):
                lines.append(f"  [{stripped[3:].upper()}]")
            # Convert # problem headers to plain bold-style markers
            elif stripped.startswith("#"):
                text = re.sub(r'^#+\s*', '', stripped)
                lines.append(f"* {text}")
            else:
                lines.append(stripped)
        lines.append("")

    # To-do items (plain text labels for EMR — no emoji)
    _EMR_BUCKET_LABELS = {
        "pre_transfer": "BEFORE TRANSFER (ICU team):",
        "ward_ongoing": "ON THE WARD (receiving team):",
        "discharge": "AT DISCHARGE (case manager/team):",
    }
    todo_items = output.get("todo_checklist", [])
    if todo_items:
        from icu_pause.rendering.formatter import _BUCKET_ORDER
        normalized = [
            t if isinstance(t, dict) else {"bucket": "ward_ongoing", "text": t}
            for t in todo_items
        ]
        lines.append("TO-DO LIST")
        lines.append("-" * 40)
        for bkey in _BUCKET_ORDER:
            items = [t["text"] for t in normalized if t.get("bucket", "ward_ongoing") == bkey]
            if not items:
                continue
            lines.append(f"  {_EMR_BUCKET_LABELS[bkey]}")
            for text in items:
                lines.append(f"  [ ] {text}")
        lines.append("")

    # Warnings — clinician-facing categories only.
    visible_warnings = render_warnings_for_clinician(output.get("warnings", []))
    if visible_warnings:
        lines.append("WARNINGS:")
        for w in visible_warnings:
            lines.append(f"  * {_format_clinician_warning_line(w)}")
        lines.append("")

    lines.append("--- AI-generated. Physician review required. ---")

    return "\n".join(lines)
